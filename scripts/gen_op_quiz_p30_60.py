"""
운전취급규정 30~60페이지 심화 문제지 (30문항) + 답지(별지).
- 문제지: 2단 다단 레이아웃, 3페이지 내 수록 (10pt)
- 답지: 별도 페이지(단단 레이아웃) + 해설
- 정답 위치: ①7 / ②8 / ③8 / ④7 (총 30) — 골고루 분포
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = "운전취급규정_심화문제_30-60페이지.docx"

# 정답 위치는 미리 배분되어 있고, options 배열의 0번이 항상 정답인 형태로 작성.
# generate() 시 ans_pos(0~3)에 맞춰 옵션 위치를 회전시킴.

QUESTIONS = [
    # 30 페이지
    {"page": 30, "ans_pos": 1,
     "q": "1~4호선 분기기에서 #10 선로전환기 양개의 경우 제한속도는?",
     "ans": "45km/h",
     "dist": ["35km/h", "55km/h", "25km/h"],
     "exp": "제103조 1항 1호: 1~4호선 #10 편개 35km/h, 양개 45km/h."},
    {"page": 30, "ans_pos": 2,
     "q": "5~8호선 분기기 #8 양개에서의 ATC 지시속도는?",
     "ans": "35km/h",
     "dist": ["25km/h", "40km/h", "60km/h"],
     "exp": "제103조 1항 2호: 5~8호선 #8 양개 제한속도 40km/h이나 ATC 지시속도는 35km/h."},
    # 31 페이지
    {"page": 31, "ans_pos": 0,
     "q": "진행수신호 현시 또는 진행수신호 생략 통보로 출발신호기를 지나 진입할 때 25km/h 운전이 종료되는 위치는?",
     "ans": "가장 바깥쪽 선로전환기까지",
     "dist": ["다음 신호기까지", "정차 위치까지", "다음 정거장까지"],
     "exp": "제104조 1항 4호 비고: 출발신호기의 경우 가장 바깥쪽 선로전환기까지 25km/h."},
    {"page": 31, "ans_pos": 3,
     "q": "진로개통표시기 불량으로 승인을 받고 25km/h 이하로 운전하는 경우, 속도 제한이 끝나는 시점은?",
     "ans": "최후부 차량이 표시기를 통과할 때까지",
     "dist": ["다음 신호기 통과까지", "정거장 진입까지", "운전관제 승인 해제까지"],
     "exp": "제104조 1항 6호 비고: 최후부 차량이 표시기를 통과할 때까지 25km/h 이하."},
    {"page": 31, "ans_pos": 1,
     "q": "1~4호선 ATC 차상장치 고장 확인을 위해 운전관제 지시로 다음 신호기까지 확인운전할 때의 속도는?",
     "ans": "15km/h",
     "dist": ["25km/h", "45km/h", "정상 속도"],
     "exp": "제104조 1항 7호 비고: ATC 차상장치 고장 확인 다음 신호기까지 15km/h."},
    {"page": 31, "ans_pos": 2,
     "q": "선로전환기에 대향하여 운전할 경우 적용되는 일반 제한속도와 예외 조건은?",
     "ans": "25km/h, 연동장치에 의해 잠금장치된 경우는 제외",
     "dist": ["45km/h, 모든 경우 적용", "15km/h, 무조건 적용", "65km/h, 정거장 통과 시 적용"],
     "exp": "제104조 1항 8호: 선로전환기 대향 운전 25km/h, 연동 잠금 시 제외."},
    {"page": 31, "ans_pos": 3,
     "q": "정거장의 승강장을 통과 운전할 때의 제한속도(정거장에 정차하는 열차는 제외)는?",
     "ans": "45km/h",
     "dist": ["25km/h", "65km/h", "35km/h"],
     "exp": "제104조 1항 9호: 승강장 통과 운전 시 45km/h, 정차 열차 예외."},
    # 32~33 페이지
    {"page": 32, "ans_pos": 0,
     "q": "1~4호선 차내신호 '25(YARD)' 신호의 비고 항목으로 옳은 것은?",
     "ans": "구내운전구간 현시",
     "dist": ["일단 정차 시 15km/h 현시", "본선 정상 운전", "정지 후 운전관제 승인"],
     "exp": "제105조 1항 2호: 25(YARD)는 구내운전구간 현시, 25km/h 운전."},
    {"page": 33, "ans_pos": 1,
     "q": "5~8호선 ATC 80신호에서 ATO 정상운전·회복운전 속도가 옳게 짝지어진 것은?",
     "ans": "정상 75km/h · 회복 77km/h",
     "dist": ["정상 80km/h · 회복 80km/h", "정상 75km/h · 회복 80km/h", "정상 70km/h · 회복 72km/h"],
     "exp": "제105조 2항 표: 80신호 — 정상 75, 회복 77."},
    {"page": 33, "ans_pos": 2,
     "q": "1000분의 10 이상의 선로에 차량을 유치할 수 없으나, 예외가 인정되는 경우는?",
     "ans": "기동되어 있는 차량인 경우",
     "dist": ["사장 승인을 받은 경우", "관제 통보만 한 경우", "비상시"],
     "exp": "제108조 2항: 1/1000의 10 이상 본선 유치 금지, 기동된 차량은 예외."},
    # 34 페이지
    {"page": 34, "ans_pos": 3,
     "q": "다음 중 차량 유치가 금지된 선로에 해당하지 않는 것은?",
     "ans": "주요한 본선",
     "dist": ["안전측선·주행시험선", "선로전환기 또는 분기부 내", "차량 접촉 한계 표지 바깥쪽"],
     "exp": "제109조 1항: 안전측선·주행시험·입환신호기 사용 구간·분기부·한계 바깥·특수시설. 주요 본선은 다른 규정."},
    {"page": 34, "ans_pos": 0,
     "q": "기동되어 있는 차량의 감시 책임자 중, 차량기지 내 입환을 제외한 그 외의 경우 책임자는?",
     "ans": "검사(검수)책임자",
     "dist": ["담당 승무원", "차장", "구내기관사"],
     "exp": "제111조 2호: 차량기지 입환은 구내기관사, 그 외는 검사(검수)책임자."},
    {"page": 34, "ans_pos": 1,
     "q": "열차를 조성하는 차량에 비치해야 할 바퀴굄목 비치 기준은?",
     "ans": "앞·뒤 차량에 각 2개 이상",
     "dist": ["앞 차량 1개", "후부 차량 2개", "모든 차량 1개 이상"],
     "exp": "제112조 2호: 앞·뒤 차량에 각 2개 이상."},
    # 35 페이지
    {"page": 35, "ans_pos": 2,
     "q": "개폐식 구름방지장치의 설치 위치는?",
     "ans": "본선 분기 측선의 차량접촉한계표지 안쪽 3m 지점",
     "dist": ["본선 분기점 외측 3m", "한계표지 바깥쪽 5m", "정거장 중앙"],
     "exp": "제114조 2항: 본선 분기 측선의 한계표지 안쪽 3m 지점 설치."},
    # 36 페이지
    {"page": 36, "ans_pos": 3,
     "q": "단선 구간의 상용폐색방식으로 시행되는 것은?",
     "ans": "단선 자동폐색식",
     "dist": ["차내신호 폐색식", "통신식", "지도통신식"],
     "exp": "제118조 1항 2호: 단선 구간 상용폐색은 단선 자동폐색식."},
    {"page": 36, "ans_pos": 0,
     "q": "차내신호 폐색식 시행 구간에서 폐색구간 경계지점은?",
     "ans": "장내경계표지·출발경계표지·폐색경계표지 설치 지점",
     "dist": ["장내·출발·폐색 신호기 설치 지점", "양방향 승강장 끝지점", "운전관제 지정 구간"],
     "exp": "제119조 3항 1호 나목: 차내신호 폐색식은 각종 경계표지 설치 지점."},
    # 37 페이지
    {"page": 37, "ans_pos": 1,
     "q": "지령식 폐색구간 경계지점(1~4호선)으로 옳은 것은?",
     "ans": "각 정거장의 열차 진행 방향 승강장의 끝지점",
     "dist": ["정거장 내외의 경계지점", "가장 바깥쪽의 장내신호기 설치 지점만", "운전관제가 지정하는 구간"],
     "exp": "제119조 3항 2호 나목: 1~4호선 지령식 — 각 정거장 진행 방향 승강장 끝지점. (운전취급역이면 가장 바깥쪽 장내신호기 또는 진로개통표시기)"},
    # 39 페이지
    {"page": 39, "ans_pos": 2,
     "q": "지도통신식 폐색의 취급자는?",
     "ans": "소장",
     "dist": ["역장", "운전관제", "운전취급자"],
     "exp": "제126조 2항 2호: 지도통신식 폐색 취급자는 소장."},
    # 40 페이지
    {"page": 40, "ans_pos": 3,
     "q": "자동폐색식 또는 차내신호폐색식 시행 구간에서 자동으로 정지신호가 현시되는 조건이 아닌 것은?",
     "ans": "기관사의 임의 정지 요청이 있을 때",
     "dist": ["폐색구간에 열차 또는 차량이 있을 때", "폐색구간 선로전환기가 정당한 방향으로 개통되지 않았을 때", "폐색장치 고장 발생 시"],
     "exp": "제130조 1항: 열차 존재·선로전환기 미개통·분기 지장·폐색 고장·반대 진행 시. 기관사 임의 요청은 사유 아님."},
    # 41 페이지
    {"page": 41, "ans_pos": 0,
     "q": "1~4호선 자동폐색 구간에서 출발신호기 사용 불가 시 폐색방식을 변경하지 않을 수 있는 조건이 아닌 것은?",
     "ans": "차장이 직접 차량 유무를 확인한 경우",
     "dist": ["열차운행표시판으로 차량 없음 확인", "입환신호기 또는 인접선로 출발신호기로 차량 없음 확인", "적임자를 파견하여 차량 없음 확인"],
     "exp": "제132조 2호: 표시판·입환신호기·적임자 확인 3가지가 변경 미요 조건. 차장 단독 확인은 인정 안 됨."},
    # 42 페이지
    {"page": 42, "ans_pos": 1,
     "q": "통신식 시행 시 운전명령서의 보존 기간은? (사고 관련 미해당)",
     "ans": "3개월",
     "dist": ["1년", "5년", "6개월"],
     "exp": "제138조 3항: 운전명령서 3개월 보존, 사고 관련은 5년."},
    {"page": 42, "ans_pos": 2,
     "q": "통신식 폐색수속에서 출발역의 '○○열차 폐색' 통보에 대한 상대역의 응답으로 옳은 것은?",
     "ans": "'○○열차 폐색승인'",
     "dist": ["'○○열차 도착'", "'○○열차 출발'", "'○○열차 진행'"],
     "exp": "제139조 1항 2호: 상대역은 '○○열차 폐색승인'으로 응답."},
    # 43 페이지
    {"page": 43, "ans_pos": 3,
     "q": "폐색구간 상태 표시판 중 '폐색표(열차 폐색구간에 있음)'의 색깔은?",
     "ans": "백색판에 적색 글씨",
     "dist": ["적색판에 백색 글씨", "황색판에 흑색 글씨", "백색판에 흑색 글씨"],
     "exp": "제141조 2항: 폐색표 — 백색판에 적색 글씨, 개통표 — 백색판에 흑색 글씨."},
    # 47 페이지
    {"page": 47, "ans_pos": 0,
     "q": "지도통신식의 지도표 발행 번호 범위와 사용 방법은?",
     "ans": "1호~10호까지 순차 발행, 10호 발행 후 1호부터 재발행",
     "dist": ["51호~100호 1열차 1매", "1호~50호 순환", "1호~99호 사용 후 폐기"],
     "exp": "제149조 3항: 지도표 1~10호 순차 발행, 10호 후 1호 재발행."},
    {"page": 47, "ans_pos": 1,
     "q": "지도권의 발행 번호 범위와 발행 방식은?",
     "ans": "51호~100호 / 1열차 1매씩 순차 발행",
     "dist": ["1호~10호 순환", "1호~50호 순차", "101호~150호 1열차 1매"],
     "exp": "제150조 3항: 지도권 51~100호, 1열차 1매씩 순차 발행."},
    # 49 페이지
    {"page": 49, "ans_pos": 2,
     "q": "폐색수속 후 사고 등으로 폐색구간에 진입시킬 수 없다고 인정할 때, 폐색수속을 일단 취소할 수 있는 시간 기준은?",
     "ans": "10분 이내",
     "dist": ["5분 이내", "30분 이내", "즉시(시간 무관)"],
     "exp": "제154조: 폐색수속 후 10분 이내 진입 불가 시 취소, 지도권은 무효(×)."},
    # 50 페이지
    {"page": 50, "ans_pos": 3,
     "q": "사고에 관련된 지도표의 보존 기간은?",
     "ans": "5년",
     "dist": ["1년", "3개월", "10년"],
     "exp": "제159조 2항: 일반 지도표 3개월 보존, 사고 관련 5년 보존."},
    # 51 페이지
    {"page": 51, "ans_pos": 0,
     "q": "전령자가 착용하는 완장의 색깔로 옳은 것은?",
     "ans": "흰 바탕에 적색 문자",
     "dist": ["적색 바탕에 흰색 문자", "황색 바탕에 흑색 문자", "녹색 바탕에 백색 문자"],
     "exp": "제164조 2항: 전령자 완장은 흰 바탕에 적색 문자, 380×90mm."},
    # 53 페이지
    {"page": 53, "ans_pos": 1,
     "q": "신호기와 수신호가 서로 다른 신호를 현시하고 있을 경우의 적용 원칙은?",
     "ans": "최대의 제한을 받는 신호 (단, 수신호에 의할 것을 통보받았을 경우는 수신호)",
     "dist": ["항상 신호기 우선", "항상 수신호 우선", "기관사가 판단"],
     "exp": "제176조 2항: 양자 다를 때 최대 제한 신호 따름, 수신호 통보 시 수신호."},
    # 54 페이지
    {"page": 54, "ans_pos": 2,
     "q": "상치신호기의 분류 중 '종속신호기'에 해당하는 것은?",
     "ans": "원방신호기·중계신호기·궤도밀착신호기",
     "dist": ["차내신호기·장내신호기·출발신호기", "진로표시기·진로개통표시기", "서행예고·서행·서행해제신호기"],
     "exp": "제180조 1항 2호: 종속신호기 — 원방·중계·궤도밀착 (주신호기 — 차내·장내·출발·폐색·입환)."},
]


# 정답 위치에 맞춰 보기 배열 회전
def options_at_position(q):
    pos = q["ans_pos"]  # 0~3
    opts = [None, None, None, None]
    opts[pos] = q["ans"]
    j = 0
    for i in range(4):
        if opts[i] is None:
            opts[i] = q["dist"][j]
            j += 1
    return opts


def set_section_columns(section, num=2):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), "360")  # 단 간격 18pt


def set_margins(section, top=1.2, bottom=1.2, left=1.5, right=1.5):
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)

    # ── 섹션 1: 문제지 (2단 레이아웃, 3페이지 내) ──
    sec1 = doc.sections[0]
    set_margins(sec1, top=1.4, bottom=1.4, left=1.4, right=1.4)
    set_section_columns(sec1, num=2)

    # 표지(2단 위 헤더 — column-spanning은 한 줄로 위치)
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("운전취급규정 심화문제 (30~60페이지)")
    r.bold = True
    r.font.size = Pt(14)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("난이도 상 · 30문항 · 시험 시간 권장 30분")
    sr.font.size = Pt(9.5)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # 문제 본문 (2단 흐름)
    for i, q in enumerate(QUESTIONS, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        rn = p.add_run(f"{i}. ")
        rn.bold = True
        rn.font.size = Pt(10)
        p.add_run(q["q"]).font.size = Pt(10)

        opts = options_at_position(q)
        for idx, opt in enumerate(opts):
            po = doc.add_paragraph()
            po.paragraph_format.left_indent = Cm(0.4)
            po.paragraph_format.space_after = Pt(0)
            po.add_run(f"  {['①','②','③','④'][idx]} {opt}").font.size = Pt(9.5)

        # 문항 간 빈 줄
        gap = doc.add_paragraph()
        gap.paragraph_format.space_after = Pt(2)

    # ── 섹션 2: 답지 (단단) ──
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(new_section, top=1.6, bottom=1.6, left=2.0, right=2.0)
    set_section_columns(new_section, num=1)

    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = h2.add_run("정답 및 해설")
    r2.bold = True
    r2.font.size = Pt(15)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ir = info.add_run("ㆍ근거 조문을 함께 정리하였습니다.")
    ir.font.size = Pt(10)
    ir.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # 정답표 (가로형)
    table = doc.add_table(rows=2, cols=10)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    ans_cells = table.rows[1].cells
    counter = 0
    grid = []  # 30/10 = 3 rows of 10
    for batch_start in range(0, 30, 10):
        if batch_start > 0:
            table.add_row()
            table.add_row()
        nums_row = table.rows[-2 if batch_start > 0 else 0]
        ans_row = table.rows[-1 if batch_start > 0 else 1]
        for c in range(10):
            qi = batch_start + c
            if qi >= 30:
                break
            nums_row.cells[c].text = str(qi + 1)
            ans_row.cells[c].text = ["①", "②", "③", "④"][QUESTIONS[qi]["ans_pos"]]
            # 폰트 정렬
            for p in nums_row.cells[c].paragraphs + ans_row.cells[c].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.bold = True

    doc.add_paragraph()

    # 해설 — 단단 (각 문항 정답 ○ + 근거)
    for i, q in enumerate(QUESTIONS, start=1):
        p = doc.add_paragraph()
        rn = p.add_run(f"{i}. ")
        rn.bold = True
        marker = ["①", "②", "③", "④"][q["ans_pos"]]
        rm = p.add_run(f"정답 {marker}  ")
        rm.bold = True
        rm.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        p.add_run(f"({q['ans']})\n해설: {q['exp']}").font.size = Pt(10)

    doc.save(OUT_PATH)
    # 정답 분포 카운트
    from collections import Counter
    dist = Counter(q["ans_pos"] for q in QUESTIONS)
    print(f"OK saved: {OUT_PATH} ({len(QUESTIONS)}문항)")
    print(f"정답 분포: ①={dist[0]}, ②={dist[1]}, ③={dist[2]}, ④={dist[3]}")


if __name__ == "__main__":
    main()
